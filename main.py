from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import Inches
import qrcode
import os

def generate_qr_code_image(data, filename):
    """
    Generate a QR code with the given data (e.g. SMS link) and save as a PNG.
    """
    img = qrcode.make(data)  # uses 'qrcode' to generate QR
    img.save(filename)

def has_placeholder(shape, placeholder_alttext):
    return shape._inline.docPr.attrib['descr'] == placeholder_alttext

def find_run(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if 'docPr' in run._r.xml:  # means that run's XML had our shape
                return run

def replace_placeholder_with_qr(input_docx, output_docx, placeholder_alttext, qr_data):
    """
    1. Open the Word document (input_docx).
    2. Find any inline shape (image) whose alt_text == placeholder_alttext.
    3. Generate a QR code (using qr_data).
    4. Delete the placeholder shape.
    5. Insert a new image (the QR code) at the same location, with the same size.
    6. Save the new document (output_docx).
    """
    # Generate the QR code image
    temp_qr_file = "temp_qr.png"
    generate_qr_code_image(qr_data, temp_qr_file)

    doc = Document(input_docx)
    # run = find_run(doc)

    for shape in doc.inline_shapes:
        # Check if this shape is our placeholder
        if not has_placeholder(shape, placeholder_alttext):
            continue

        # 1) Get the current shape's dimensions
        old_width = shape.width
        old_height = shape.height

        # 2) Remove the old placeholder shape from the document XML
        inline = shape._inline  # the actual XML element
        drawing_elm = inline.getparent()
        run_elm = drawing_elm.getparent()
        paragraph_elm = run_elm.getparent()

        run_elm.remove(drawing_elm)

        paragraph = Paragraph(paragraph_elm, doc._body)
        run = Run(run_elm, paragraph)
        run.add_picture(temp_qr_file, width=old_width, height=old_height)


    # Save the modified doc
    doc.save(output_docx)

    # Clean up temporary QR file
    if os.path.exists(temp_qr_file):
        os.remove(temp_qr_file)

def main():
    input_docx = "template.docx"              # Your DOCX with placeholder image
    output_docx = "output_with_qr.docx"       # Where to save the result
    placeholder_alttext = "QRCODE_PLACEHOLDER"
    # The data for an "SMS" QR code typically looks like:
    # "sms:+1234567890?body=Hello" or "SMSTO:+1234567890:Hello"
    qr_data = "sms:+15551234567?body=HelloWorld"

    replace_placeholder_with_qr(input_docx, output_docx, placeholder_alttext, qr_data)
    print(f"Created {output_docx} with QR code replacing the placeholder image.")

if __name__ == "__main__":
    main()
